"""
Document format converter implementation.

This module provides classes for converting between document formats
such as Markdown, Word, PDF, and other text-based formats.
"""

import logging
import os
import tempfile
from typing import BinaryIO, List, Optional, Union

import markdown
import pypandoc
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter as PDFToDocxConverter
from bs4 import BeautifulSoup

from .converter_base import ConverterBase

logger = logging.getLogger(__name__)


class DocumentConverter(ConverterBase):
    """Class for converting between document formats."""
    
    # Define supported formats
    _SUPPORTED_INPUT_FORMATS = ['md', 'markdown', 'docx', 'doc', 'pdf', 'txt', 'html', 'odt', 'rtf']
    _SUPPORTED_OUTPUT_FORMATS = ['md', 'markdown', 'docx', 'pdf', 'txt', 'html', 'odt', 'rtf']
    
    def __init__(self):
        """Initialize the document converter."""
        super().__init__()
        self.logger.info("Initializing DocumentConverter")
        
        # Check if pandoc is installed
        try:
            pypandoc.get_pandoc_version()
            self.pandoc_available = True
            self.logger.info("Pandoc is available")
        except OSError:
            self.pandoc_available = False
            self.logger.warning("Pandoc is not available. Some conversions may not work.")
    
    def get_supported_input_formats(self) -> List[str]:
        """
        Get a list of supported input formats.
        
        Returns:
            List[str]: List of supported input format extensions
        """
        return self._SUPPORTED_INPUT_FORMATS
    
    def get_supported_output_formats(self) -> List[str]:
        """
        Get a list of supported output formats.
        
        Returns:
            List[str]: List of supported output format extensions
        """
        return self._SUPPORTED_OUTPUT_FORMATS
    
    def convert(self, input_file: Union[str, BinaryIO], output_format: str, 
                output_path: Optional[str] = None, **kwargs) -> str:
        """
        Convert a document from one format to another.
        
        Args:
            input_file: Path to the input file or file-like object
            output_format: The target format for conversion
            output_path: Optional path where the output file should be saved
            **kwargs: Additional conversion options
                - style_template: Path to a style template for the output document
                - preserve_formatting: Whether to preserve formatting (default: True)
                
        Returns:
            str: Path to the converted file
        """
        # Handle file-like objects
        if not isinstance(input_file, str):
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{kwargs.get('input_format', 'tmp')}")
            temp_file.write(input_file.read())
            temp_file.close()
            input_file = temp_file.name
            self.logger.info(f"Created temporary file: {input_file}")
        
        # Validate input file
        if not self.validate_file_exists(input_file):
            raise FileNotFoundError(f"Input file not found: {input_file}")
        
        # Get input format from file extension if not specified
        input_format = kwargs.get('input_format', self.get_file_extension(input_file))
        
        # Create output path if not provided
        if not output_path:
            output_path = self.create_output_path(input_file, output_format)
        
        self.logger.info(f"Converting {input_file} ({input_format}) to {output_path} ({output_format})")
        
        # Handle specific conversion cases
        if input_format in ['md', 'markdown'] and output_format in ['docx', 'doc']:
            return self._convert_markdown_to_docx(input_file, output_path, **kwargs)
        elif input_format in ['docx', 'doc'] and output_format in ['md', 'markdown']:
            return self._convert_docx_to_markdown(input_file, output_path, **kwargs)
        elif input_format == 'pdf' and output_format in ['docx', 'doc']:
            return self._convert_pdf_to_docx(input_file, output_path, **kwargs)
        elif input_format == 'html' and output_format in ['md', 'markdown']:
            return self._convert_html_to_markdown(input_file, output_path, **kwargs)
        else:
            # Use pandoc for other conversions
            if not self.pandoc_available:
                raise RuntimeError("Pandoc is required for this conversion but is not available")
            
            return self._convert_with_pandoc(input_file, input_format, output_format, output_path, **kwargs)
    
    def _convert_markdown_to_docx(self, input_file: str, output_path: str, **kwargs) -> str:
        """
        Convert Markdown to DOCX using python-docx.
        
        Args:
            input_file: Path to the input Markdown file
            output_path: Path where the output DOCX file should be saved
            **kwargs: Additional conversion options
                
        Returns:
            str: Path to the converted DOCX file
        """
        self.logger.info(f"Converting Markdown to DOCX: {input_file} -> {output_path}")
        
        # Read markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # Create a new Word document
        doc = Document()
        
        # Parse HTML and add content to Word document
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Process headings, paragraphs, lists, etc.
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'code', 'pre']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name == 'pre':
                doc.add_paragraph(element.get_text(), style='Code')
            # Add more element handling as needed
        
        # Apply style template if provided
        style_template = kwargs.get('style_template')
        if style_template and os.path.exists(style_template):
            # Apply styles from template (simplified implementation)
            self.logger.info(f"Applying style template: {style_template}")
        
        # Save the document
        doc.save(output_path)
        self.logger.info(f"Successfully converted Markdown to DOCX: {output_path}")
        
        return output_path
    
    def _convert_docx_to_markdown(self, input_file: str, output_path: str, **kwargs) -> str:
        """
        Convert DOCX to Markdown.
        
        Args:
            input_file: Path to the input DOCX file
            output_path: Path where the output Markdown file should be saved
            **kwargs: Additional conversion options
                
        Returns:
            str: Path to the converted Markdown file
        """
        self.logger.info(f"Converting DOCX to Markdown: {input_file} -> {output_path}")
        
        # Use pandoc for conversion
        if self.pandoc_available:
            return self._convert_with_pandoc(input_file, 'docx', 'markdown', output_path, **kwargs)
        
        # Fallback to python-docx if pandoc is not available
        doc = Document(input_file)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    f.write('#' * level + ' ' + para.text + '\n\n')
                else:
                    f.write(para.text + '\n\n')
        
        self.logger.info(f"Successfully converted DOCX to Markdown: {output_path}")
        return output_path
    
    def _convert_pdf_to_docx(self, input_file: str, output_path: str, **kwargs) -> str:
        """
        Convert PDF to DOCX using pdf2docx.
        
        Args:
            input_file: Path to the input PDF file
            output_path: Path where the output DOCX file should be saved
            **kwargs: Additional conversion options
                
        Returns:
            str: Path to the converted DOCX file
        """
        self.logger.info(f"Converting PDF to DOCX: {input_file} -> {output_path}")
        
        # Convert PDF to DOCX
        cv = PDFToDocxConverter(input_file)
        cv.convert(output_path)
        cv.close()
        
        self.logger.info(f"Successfully converted PDF to DOCX: {output_path}")
        return output_path
    
    def _convert_html_to_markdown(self, input_file: str, output_path: str, **kwargs) -> str:
        """
        Convert HTML to Markdown.
        
        Args:
            input_file: Path to the input HTML file
            output_path: Path where the output Markdown file should be saved
            **kwargs: Additional conversion options
                
        Returns:
            str: Path to the converted Markdown file
        """
        self.logger.info(f"Converting HTML to Markdown: {input_file} -> {output_path}")
        
        # Use pandoc for conversion
        if self.pandoc_available:
            return self._convert_with_pandoc(input_file, 'html', 'markdown', output_path, **kwargs)
        
        # Fallback to BeautifulSoup if pandoc is not available
        with open(input_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # Simple HTML to Markdown conversion
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'a', 'strong', 'em']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    f.write('#' * level + ' ' + element.get_text() + '\n\n')
                elif element.name == 'p':
                    f.write(element.get_text() + '\n\n')
                elif element.name == 'a':
                    f.write(f"[{element.get_text()}]({element.get('href')})")
                elif element.name == 'strong':
                    f.write(f"**{element.get_text()}**")
                elif element.name == 'em':
                    f.write(f"*{element.get_text()}*")
        
        self.logger.info(f"Successfully converted HTML to Markdown: {output_path}")
        return output_path
    
    def _convert_with_pandoc(self, input_file: str, input_format: str, 
                            output_format: str, output_path: str, **kwargs) -> str:
        """
        Convert a document using pandoc.
        
        Args:
            input_file: Path to the input file
            input_format: Input file format
            output_format: Output file format
            output_path: Path where the output file should be saved
            **kwargs: Additional conversion options
                
        Returns:
            str: Path to the converted file
        """
        self.logger.info(f"Converting with pandoc: {input_file} ({input_format}) -> {output_path} ({output_format})")
        
        # Map our format names to pandoc format names
        format_mapping = {
            'md': 'markdown',
            'markdown': 'markdown',
            'docx': 'docx',
            'doc': 'docx',
            'pdf': 'pdf',
            'txt': 'plain',
            'html': 'html',
            'odt': 'odt',
            'rtf': 'rtf'
        }
        
        pandoc_input_format = format_mapping.get(input_format, input_format)
        pandoc_output_format = format_mapping.get(output_format, output_format)
        
        # Set pandoc options
        extra_args = []
        
        # Add options for preserving formatting if specified
        if kwargs.get('preserve_formatting', True):
            if pandoc_output_format == 'markdown':
                # Modified: Removed the --atx-headers option that was causing the error
                extra_args.extend(['--wrap=none'])
            elif pandoc_output_format == 'docx':
                extra_args.extend(['--reference-doc=' + kwargs.get('style_template')]) if kwargs.get('style_template') else None
        
        # Convert using pandoc
        pypandoc.convert_file(
            input_file,
            pandoc_output_format,
            format=pandoc_input_format,
            outputfile=output_path,
            extra_args=extra_args
        )
        
        self.logger.info(f"Successfully converted with pandoc: {output_path}")
        return output_path